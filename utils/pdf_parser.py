"""
Document parsing utilities for CV/Resume extraction.
Supports PDF, Word (.docx), and plain text files.
"""

import io
from typing import Tuple
from Pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from a PDF file.
    
    Args:
        file_bytes: Raw bytes of the PDF file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = "\n".join(text_parts).strip()
        
        if not full_text:
            return "", "The PDF appears to be empty or contains only images. Please upload a PDF with selectable text."
        
        return full_text, ""
        
    except Exception as e:
        return "", f"Failed to parse PDF: {str(e)}. Please try uploading a different file."


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from a Word document (.docx).
    
    Args:
        file_bytes: Raw bytes of the Word file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts).strip()
        
        if not full_text:
            return "", "The Word document appears to be empty. Please upload a document with text content."
        
        return full_text, ""
        
    except Exception as e:
        return "", f"Failed to parse Word document: {str(e)}. Please try uploading a different file."


def extract_text_from_txt(file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from a plain text file.
    
    Args:
        file_bytes: Raw bytes of the text file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        text = file_bytes.decode('utf-8').strip()
        
        if not text:
            return "", "The text file appears to be empty. Please upload a file with content."
        
        return text, ""
        
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode('latin-1').strip()
            if not text:
                return "", "The text file appears to be empty."
            return text, ""
        except Exception as e:
            return "", f"Failed to read text file: {str(e)}"
    except Exception as e:
        return "", f"Failed to read text file: {str(e)}"


def parse_document(uploaded_file) -> Tuple[str, str]:
    """
    Parse an uploaded document and extract text.
    Automatically detects file type and uses appropriate parser.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    if uploaded_file is None:
        return "", "No file uploaded."
    
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()
    
    if file_name.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif file_name.endswith('.txt'):
        return extract_text_from_txt(file_bytes)
    else:
        return "", "Unsupported file format. Please upload a PDF, Word (.docx), or text (.txt) file."
